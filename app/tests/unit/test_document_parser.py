from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document

from app.services.parsing.document_parser import DocumentParser



def test_extract_pdf_text(tmp_path: Path) -> None:
    path = tmp_path / "resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Alex Johnson\nData Engineer")
    doc.save(path)
    doc.close()

    parser = DocumentParser()
    text = parser.extract_text(path)

    assert "Alex Johnson" in text



def test_extract_docx_text(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Jamie Lee")
    document.add_paragraph("ML Engineer")
    document.save(path)

    parser = DocumentParser()
    text = parser.extract_text(path)

    assert "Jamie Lee" in text
